from __future__ import annotations

import io
import subprocess
import tempfile
from html import escape
from pathlib import Path
from typing import Any

from bs4 import BeautifulSoup
from markdownify import markdownify

SUPPORTED_EXTENSIONS = {
    ".html",
    ".htm",
    ".xhtml",
    ".pdf",
    ".doc",
    ".docx",
    ".odt",
    ".rtf",
    ".xls",
    ".xlsx",
}


def parse_document(
    data: bytes, filename: str, *, mode: str = "auto", max_pages: int | None = None
) -> tuple[str, dict[str, Any]]:
    suffix = Path(filename).suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"unsupported document type: {suffix or 'unknown'}")
    if suffix in {".html", ".htm", ".xhtml"}:
        return _parse_html(data), {"fileType": "html"}
    if suffix == ".pdf":
        return _parse_pdf(data, mode=mode, max_pages=max_pages)
    if suffix == ".docx":
        return _parse_docx(data), {"fileType": "docx"}
    if suffix == ".doc":
        return _parse_doc(data), {"fileType": "doc"}
    if suffix == ".odt":
        return _parse_odt(data), {"fileType": "odt"}
    if suffix == ".rtf":
        return _parse_rtf(data), {"fileType": "rtf"}
    return _parse_workbook(data, suffix), {"fileType": suffix.removeprefix(".")}


def rendered_formats(markdown: str, formats: list[str]) -> dict[str, Any]:
    result: dict[str, Any] = {}
    if "markdown" in formats:
        result["markdown"] = markdown
    if "html" in formats or "rawHtml" in formats:
        html = "\n".join(f"<p>{escape(line)}</p>" for line in markdown.splitlines() if line)
        if "html" in formats:
            result["html"] = html
        if "rawHtml" in formats:
            result["rawHtml"] = html
    if "summary" in formats:
        paragraphs = [part.strip() for part in markdown.split("\n\n") if part.strip()]
        result["summary"] = paragraphs[0][:1000] if paragraphs else ""
    if "json" in formats:
        result["json"] = {"text": markdown}
    return result


def _parse_html(data: bytes) -> str:
    soup = BeautifulSoup(data, "html.parser")
    for node in soup(["script", "style", "noscript"]):
        node.decompose()
    return markdownify(str(soup), heading_style="ATX").strip()


def _parse_pdf(data: bytes, *, mode: str, max_pages: int | None) -> tuple[str, dict[str, Any]]:
    import fitz

    document = fitz.open(stream=data, filetype="pdf")
    total_pages = len(document)
    page_count = min(max_pages or total_pages, total_pages)
    pages: list[str] = []
    ocr_pages = 0
    for page in list(document)[:page_count]:
        text = page.get_text("text").strip()
        if mode == "ocr" or (mode == "auto" and len(text) < 20):
            text = _ocr_page(page)
            ocr_pages += 1
        pages.append(text)
    return "\n\n".join(page for page in pages if page).strip(), {
        "fileType": "pdf",
        "numPages": page_count,
        "totalPages": total_pages,
        "ocrPages": ocr_pages,
    }


def _ocr_page(page) -> str:
    import pytesseract
    from PIL import Image

    pixmap = page.get_pixmap(dpi=200, alpha=False)
    return pytesseract.image_to_string(Image.open(io.BytesIO(pixmap.tobytes("png")))).strip()


def _parse_docx(data: bytes) -> str:
    from docx import Document

    document = Document(io.BytesIO(data))
    parts = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        parts.append(
            _markdown_table([[cell.text.strip() for cell in row.cells] for row in table.rows])
        )
    return "\n\n".join(part for part in parts if part)


def _parse_doc(data: bytes) -> str:
    with tempfile.NamedTemporaryFile(suffix=".doc") as source:
        source.write(data)
        source.flush()
        result = subprocess.run(
            ["antiword", source.name], capture_output=True, check=False, timeout=60
        )
    if result.returncode != 0:
        raise ValueError(result.stderr.decode("utf-8", errors="replace").strip())
    return result.stdout.decode("utf-8", errors="replace").strip()


def _parse_odt(data: bytes) -> str:
    from odf import teletype
    from odf.opendocument import load
    from odf.text import H, P

    document = load(io.BytesIO(data))
    parts = []
    for node in document.getElementsByType(H) + document.getElementsByType(P):
        text = teletype.extractText(node).strip()
        if text:
            parts.append(text)
    return "\n\n".join(parts)


def _parse_rtf(data: bytes) -> str:
    from striprtf.striprtf import rtf_to_text

    return rtf_to_text(data.decode("latin-1", errors="replace")).strip()


def _parse_workbook(data: bytes, suffix: str) -> str:
    if suffix == ".xls":
        import xlrd

        workbook = xlrd.open_workbook(file_contents=data)
        sheets = [
            (
                sheet.name,
                [
                    [sheet.cell_value(row, col) for col in range(sheet.ncols)]
                    for row in range(sheet.nrows)
                ],
            )
            for sheet in workbook.sheets()
        ]
    else:
        from openpyxl import load_workbook

        workbook = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        sheets = [
            (sheet.title, [list(row) for row in sheet.iter_rows(values_only=True)])
            for sheet in workbook
        ]
    parts = []
    for name, rows in sheets:
        normalized = [["" if value is None else str(value) for value in row] for row in rows]
        parts.extend([f"## {name}", _markdown_table(normalized)])
    return "\n\n".join(part for part in parts if part)


def _markdown_table(rows: list[list[Any]]) -> str:
    if not rows:
        return ""
    width = max(len(row) for row in rows)
    padded = [[str(value) for value in row] + [""] * (width - len(row)) for row in rows]
    lines = ["| " + " | ".join(padded[0]) + " |", "| " + " | ".join(["---"] * width) + " |"]
    lines.extend("| " + " | ".join(row) + " |" for row in padded[1:])
    return "\n".join(lines)
