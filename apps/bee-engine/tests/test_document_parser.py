import io
import base64

from fastapi.testclient import TestClient

from bee_engine import document_parser
from bee_engine.app import app


def test_parses_html_rtf_docx_and_xlsx_uploads() -> None:
    from docx import Document
    from openpyxl import Workbook

    html, html_meta = document_parser.parse_document(
        b"<html><body><h1>Upload title</h1><p>Useful text</p><script>bad()</script></body></html>",
        "upload.html",
    )
    assert "# Upload title" in html
    assert "bad()" not in html
    assert html_meta["fileType"] == "html"

    rtf, _ = document_parser.parse_document(b"{\\rtf1\\ansi Hello \\b BeeCrawl\\b0}", "note.rtf")
    assert "Hello BeeCrawl" in rtf

    document = Document()
    document.add_heading("DOCX title", level=1)
    document.add_paragraph("DOCX body")
    docx_buffer = io.BytesIO()
    document.save(docx_buffer)
    docx, _ = document_parser.parse_document(docx_buffer.getvalue(), "report.docx")
    assert "DOCX title" in docx and "DOCX body" in docx

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Metrics"
    sheet.append(["Name", "Value"])
    sheet.append(["pages", 42])
    xlsx_buffer = io.BytesIO()
    workbook.save(xlsx_buffer)
    xlsx, _ = document_parser.parse_document(xlsx_buffer.getvalue(), "metrics.xlsx")
    assert "## Metrics" in xlsx
    assert "| pages | 42 |" in xlsx


def test_scanned_pdf_auto_mode_uses_ocr(monkeypatch) -> None:
    import fitz

    document = fitz.open()
    document.new_page()
    data = document.tobytes()
    monkeypatch.setattr(document_parser, "_ocr_page", lambda _page: "OCR recovered text")
    markdown, metadata = document_parser.parse_document(data, "scan.pdf", mode="auto")
    assert markdown == "OCR recovered text"
    assert metadata["ocrPages"] == 1


def test_applicable_parse_formats_are_rendered() -> None:
    result = document_parser.rendered_formats(
        "First paragraph\n\nSecond paragraph", ["markdown", "html", "rawHtml", "summary", "json"]
    )
    assert result["markdown"].startswith("First")
    assert result["html"].startswith("<p>")
    assert result["rawHtml"] == result["html"]
    assert result["summary"] == "First paragraph"
    assert result["json"]["text"].endswith("Second paragraph")


def test_document_parse_endpoint_returns_requested_formats() -> None:
    response = TestClient(app).post(
        "/parse",
        json={
            "filename": "upload.html",
            "base64": base64.b64encode(b"<h1>Endpoint title</h1><p>Body</p>").decode(),
            "formats": ["markdown", "html", "summary", "json"],
        },
    )
    assert response.status_code == 200
    body = response.json()
    assert "# Endpoint title" in body["data"]["markdown"]
    assert body["data"]["summary"].startswith("# Endpoint title")
    assert body["metadata"]["fileType"] == "html"
